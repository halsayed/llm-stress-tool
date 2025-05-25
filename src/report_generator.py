import os
from typing import Dict, List, Any, Optional
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches


class ReportGenerator:
    """
    Generates Word reports with charts from test results.
    """
    
    def __init__(self, results: Dict[str, Any], output_path: str):
        """
        Initialize the report generator.
        
        Args:
            results: Dictionary with test results
            output_path: Path to save the Word document
        """
        self.results = results
        self.output_path = output_path
        self.document = Document()
        self.charts_dir = os.path.dirname(output_path) + "/charts"
        
        # Create charts directory if it doesn't exist
        if not os.path.exists(self.charts_dir):
            os.makedirs(self.charts_dir)

    @staticmethod
    def _sanitize_name(name: str) -> str:
        """Sanitize a string so it can be safely used as a filename."""
        return "".join(c if c.isalnum() or c in ("-", "_") else "_" for c in name)
    
    def generate_report(self) -> None:
        """
        Generate the Word report with charts.
        """
        # Add title
        self.document.add_heading('LLM Performance Benchmark Report', 0)
        
        # Add introduction
        self.document.add_heading('Introduction', level=1)
        self.document.add_paragraph(
            'This report presents the results of performance benchmarking for Large Language Models (LLMs). '
            'The benchmarks measure tokens per second, latency, and throughput under various concurrency levels.'
        )
        
        # Add test configuration section
        self.document.add_heading('Test Configuration', level=1)
        
        # Models information
        self.document.add_heading('Models', level=2)
        models_table = self.document.add_table(rows=1, cols=2)
        models_table.style = 'Table Grid'
        header_cells = models_table.rows[0].cells
        header_cells[0].text = 'Model Name'
        header_cells[1].text = 'Base URL'
        
        for model in self.results['models']:
            row_cells = models_table.add_row().cells
            row_cells[0].text = model['model_name']
            row_cells[1].text = model['base_url']
        
        # Tests information
        self.document.add_heading('Test Cases', level=2)
        tests_table = self.document.add_table(rows=1, cols=2)
        tests_table.style = 'Table Grid'
        header_cells = tests_table.rows[0].cells
        header_cells[0].text = 'Test Name'
        header_cells[1].text = 'Input Prompt'
        
        for test in self.results['tests']:
            row_cells = tests_table.add_row().cells
            row_cells[0].text = test['name']
            row_cells[1].text = test['input']
        
        # Test parameters
        self.document.add_heading('Test Parameters', level=2)
        self.document.add_paragraph(f"Concurrency Levels: {', '.join(map(str, self.results['concurrency_levels']))}")
        self.document.add_paragraph(f"Total Requests per Test: {self.results['total_requests']}")
        
        # Generate and add charts for each model
        self.document.add_heading('Performance Results', level=1)
        
        for model_result in self.results['results']:
            model_name = model_result['model_name']
            self.document.add_heading(f"Model: {model_name}", level=2)
            
            # Generate charts for this model
            self._generate_model_charts(model_result)
            
            # Add summary tables for each test
            for test_result in model_result['test_results']:
                test_name = test_result['test_name']
                self.document.add_heading(f"Test: {test_name}", level=3)
                
                # Add summary table
                summary_table = self.document.add_table(rows=1, cols=4)
                summary_table.style = 'Table Grid'
                header_cells = summary_table.rows[0].cells
                header_cells[0].text = 'Concurrency'
                header_cells[1].text = 'Avg. Latency (s)'
                header_cells[2].text = 'Tokens/s'
                header_cells[3].text = 'Success Rate'
                
                for concurrency_result in test_result['concurrency_results']:
                    concurrency = concurrency_result['concurrency']
                    summary = concurrency_result['summary']
                    
                    row_cells = summary_table.add_row().cells
                    row_cells[0].text = str(concurrency)
                    row_cells[1].text = f"{summary['avg_latency']:.3f}"
                    row_cells[2].text = f"{summary['avg_tokens_per_second']:.2f}"
                    row_cells[3].text = f"{summary['success_rate'] * 100:.1f}%"
        
        # Save the document
        self.document.save(self.output_path)
    
    def _generate_model_charts(self, model_result: Dict[str, Any]) -> None:
        """
        Generate charts for a specific model and add them to the report.
        
        Args:
            model_result: Dictionary with model test results
        """
        model_name = model_result['model_name']
        safe_model_name = self._sanitize_name(model_name)
        
        # Prepare data for charts
        latency_data = []
        tokens_per_second_data = []
        
        for test_result in model_result['test_results']:
            test_name = test_result['test_name']
            
            for concurrency_result in test_result['concurrency_results']:
                concurrency = concurrency_result['concurrency']
                summary = concurrency_result['summary']
                
                latency_data.append({
                    'Test': test_name,
                    'Concurrency': concurrency,
                    'Latency': summary['avg_latency']
                })
                
                tokens_per_second_data.append({
                    'Test': test_name,
                    'Concurrency': concurrency,
                    'Tokens/s': summary['avg_tokens_per_second']
                })
        
        # Convert to DataFrames
        latency_df = pd.DataFrame(latency_data)
        tokens_df = pd.DataFrame(tokens_per_second_data)
        
        # Generate latency chart
        self._create_chart(
            latency_df, 
            'Concurrency', 
            'Latency', 
            'Test',
            f'Latency by Concurrency - {model_name}',
            f'{self.charts_dir}/latency_{safe_model_name}.png'
        )
        
        # Add latency chart to document
        self.document.add_heading(f"Latency by Concurrency", level=3)
        self.document.add_picture(
            f'{self.charts_dir}/latency_{safe_model_name}.png',
            width=Inches(6.0)
        )
        
        # Generate tokens per second chart
        self._create_chart(
            tokens_df, 
            'Concurrency', 
            'Tokens/s', 
            'Test',
            f'Tokens per Second by Concurrency - {model_name}',
            f'{self.charts_dir}/tokens_{safe_model_name}.png'
        )
        
        # Add tokens per second chart to document
        self.document.add_heading(f"Tokens per Second by Concurrency", level=3)
        self.document.add_picture(
            f'{self.charts_dir}/tokens_{safe_model_name}.png',
            width=Inches(6.0)
        )
    
    def _create_chart(self, df: pd.DataFrame, x_col: str, y_col: str, hue_col: str, 
                     title: str, save_path: str) -> None:
        """
        Create and save a chart.
        
        Args:
            df: DataFrame with the data
            x_col: Column name for x-axis
            y_col: Column name for y-axis
            hue_col: Column name for grouping
            title: Chart title
            save_path: Path to save the chart
        """
        plt.style.use("ggplot")
        plt.figure(figsize=(10, 6))

        colors = plt.cm.tab10.colors
        markers = ['o', 's', '^', 'D', 'v', 'P', '*', 'X', '<', '>']

        # Create line plot with enhanced styling
        for idx, (test_name, group) in enumerate(df.groupby(hue_col)):
            color = colors[idx % len(colors)]
            marker = markers[idx % len(markers)]
            plt.plot(
                group[x_col],
                group[y_col],
                marker=marker,
                linewidth=2,
                label=test_name,
                color=color,
            )

        plt.title(title, fontsize=14, fontweight="bold")
        plt.xlabel(x_col, fontsize=12)
        plt.ylabel(y_col, fontsize=12)
        plt.grid(True, linestyle="--", linewidth=0.5, alpha=0.7)
        plt.legend(frameon=False)
        plt.tight_layout()
        
        # Save the chart
        plt.savefig(save_path, dpi=300, bbox_inches='tight')
        plt.close()
    
    def add_analysis(self, analysis: Dict[str, Any]) -> None:
        """
        Add automated analysis to the report.
        
        Args:
            analysis: Dictionary with analysis results
        """
        self.document.add_heading('Automated Analysis', level=1)
        
        # Add general observations
        self.document.add_heading('General Observations', level=2)
        for observation in analysis.get('general_observations', []):
            self.document.add_paragraph(observation, style='List Bullet')
        
        # Add model comparisons if multiple models
        if len(self.results['models']) > 1:
            self.document.add_heading('Model Comparisons', level=2)
            for comparison in analysis.get('model_comparisons', []):
                self.document.add_paragraph(comparison, style='List Bullet')
        
        # Add concurrency impact analysis
        self.document.add_heading('Concurrency Impact', level=2)
        for impact in analysis.get('concurrency_impact', []):
            self.document.add_paragraph(impact, style='List Bullet')
        
        # Add recommendations
        self.document.add_heading('Recommendations', level=2)
        for recommendation in analysis.get('recommendations', []):
            self.document.add_paragraph(recommendation, style='List Bullet')
        
        # Save the updated document
        self.document.save(self.output_path)
